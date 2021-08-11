from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class Documento:
    def __init__(self, name):
        self.document = Document()
        self.name = name
        self.document.save(name)
        self.create_head_style()

    def create_head_style(self):
        my_styles = self.document.styles

        # creating the paragraph style
        self.p_style = my_styles.add_style('head-style',
                                           WD_STYLE_TYPE.PARAGRAPH)
        self.p_style.base_style = my_styles['Normal']

        self.p_style.paragraph_format.space_before = Pt(0)
        self.p_style.paragraph_format.space_after = Pt(20)
        self.p_style.font.name = 'Times New Roman'
        self.p_style.font.size = Pt(25)

    def write_page(self, title, content):
        # self.document.add_heading(title, 0)
        table = self.document.add_table(rows = 1,cols = 2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        table.cell(0,0).merge(table.cell(0,1))
        table.cell(0, 0).text = title

        par = table.cell(0,0).paragraphs[0]

        par.style = self.p_style

        par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # WD_ALIGN_PARAGRAPH.CENTER

        # defaults to True, so have to explicitly turn it off
        table.autofit = False
        table.columns[0].width = Cm(1.4)
        table.columns[1].width = Cm(14)

        for rows in content:
            row_cells = table.add_row().cells
            row_cells[0].text = str(rows[0]) + '.-'
            row_cells[1].text = str(rows[1])

        self.document.add_page_break()

        self.document.save(self.name)


if __name__ == "__main__":
    indice = Documento('PruebaIndice.docx')
    arr = []
    a = "a"*3*103

    for i in range(1,10):
        arr.append([i,a])

    indice.write_page('A',arr)
    indice.write_page('B',[[5, 6],[7, 8]])
